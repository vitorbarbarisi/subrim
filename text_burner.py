#!/usr/bin/env python3
"""
Text Burner - Processa PDFs com texto chinês e gera versão com pinyin e traduções

Usage: python3 text_burner.py <directory_name>
Example: python3 text_burner.py texto

O script:
1. Procura por arquivo PDF no diretório assets/<directory_name>
2. Extrai linhas contendo caracteres chineses do PDF
3. Cria base.txt com essas linhas
4. Chama LLM para gerar pares de tradução para cada linha
5. Sanitiza usando word-api (remove palavras com confidence_level == 3)
6. Gera novo PDF com pinyin e traduções sobrepostas verticalmente
7. É idempotente - continua de onde parou se base.txt já existir

Formato do base.txt:
- Uma linha por linha do PDF que contém caracteres chineses
- Formato: linha_original\tpares_json

Formato dos pares:
- Array JSON: ["palavra (pinyin): tradução", ...]

Exemplo de saída no PDF:
nǐ hǎo
你好
olá
"""

import sys
import argparse
import os
import json
import re
import time
from pathlib import Path
from urllib import request as urlrequest, error as urlerror
import requests

# Importar funções do processor.py
from processor import (
    _get_api_provider, _retry_api_call, _call_maritaca_pairs, _call_deepseek_pairs,
    load_dotenv
)

# Importar funções do sanitize_base.py
from sanitize_base import (
    check_word_api_health, get_word_from_api, post_word_to_api,
    extract_pairs_from_translation, process_word_api_integration
)

# Load .env immediately when module is imported
load_dotenv()

# URL da word-api localhost
WORD_API_BASE_URL = "http://localhost:7998/word-api"


def extract_text_from_pdf(pdf_path: Path) -> list[str]:
    """
    Extrai texto de um arquivo PDF.
    
    Args:
        pdf_path: Caminho para o arquivo PDF
        
    Returns:
        list: Lista de linhas do texto extraído
    """
    # Tenta PyPDF2 primeiro
    try:
        import PyPDF2
        
        lines = []
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text:
                    page_lines = text.split('\n')
                    for line in page_lines:
                        line = line.strip()
                        if line:
                            lines.append(line)
        return lines
        
    except ImportError:
        print("⚠️  PyPDF2 não encontrado, tentando pdftotext...")
        
        # Fallback para pdftotext (ferramenta do sistema)
        try:
            import subprocess
            result = subprocess.run(['pdftotext', str(pdf_path), '-'], 
                                 capture_output=True, text=True, check=True)
            lines = [line.strip() for line in result.stdout.split('\n') if line.strip()]
            return lines
        except (subprocess.CalledProcessError, FileNotFoundError):
            print("❌ Nem PyPDF2 nem pdftotext disponíveis.")
            print("Instale PyPDF2: pip install PyPDF2")
            print("Ou instale poppler-utils: brew install poppler (macOS) ou apt-get install poppler-utils (Linux)")
            sys.exit(1)
    except Exception as e:
        print(f"❌ Erro ao extrair texto do PDF: {e}")
        sys.exit(1)


def contains_chinese_characters(text: str) -> bool:
    """
    Verifica se o texto contém caracteres chineses.
    
    Args:
        text: Texto para verificar
        
    Returns:
        bool: True se contém caracteres chineses
    """
    # Range de caracteres chineses: CJK Unified Ideographs
    chinese_pattern = re.compile(r'[\u4e00-\u9fff]')
    return bool(chinese_pattern.search(text))


def create_base_file(pdf_path: Path, base_path: Path, resume: bool = False) -> bool:
    """
    Cria o arquivo base.txt com linhas contendo caracteres chineses.
    
    Args:
        pdf_path: Caminho para o arquivo PDF
        base_path: Caminho para o arquivo base.txt
        resume: Se True, continua de onde parou se base.txt já existir
        
    Returns:
        bool: True se sucesso
    """
    print(f"📖 Extraindo texto do PDF: {pdf_path.name}")
    
    # Extrai texto do PDF
    all_lines = extract_text_from_pdf(pdf_path)
    print(f"📊 Total de linhas extraídas: {len(all_lines)}")
    
    # Filtra linhas com caracteres chineses
    chinese_lines = [line for line in all_lines if contains_chinese_characters(line)]
    print(f"🔤 Linhas com caracteres chineses: {len(chinese_lines)}")
    
    if not chinese_lines:
        print("❌ Nenhuma linha com caracteres chineses encontrada no PDF")
        return False
    
    # Verifica se deve continuar de onde parou
    existing_lines = []
    if resume and base_path.exists():
        try:
            with open(base_path, 'r', encoding='utf-8') as f:
                existing_content = f.read().strip()
                if existing_content:
                    existing_lines = existing_content.split('\n')
                    print(f"📝 Arquivo base existente encontrado com {len(existing_lines)} linhas")
        except Exception as e:
            print(f"⚠️  Erro ao ler arquivo base existente: {e}")
    
    # Abre arquivo para escrita (append se resumindo)
    mode = 'a' if resume and existing_lines else 'w'
    
    with open(base_path, mode, encoding='utf-8') as f:
        # Se não está resumindo, escreve todas as linhas
        if not resume or not existing_lines:
            for line in chinese_lines:
                f.write(f"{line}\n")
            print(f"✅ Arquivo base criado com {len(chinese_lines)} linhas")
        else:
            # Se está resumindo, só adiciona linhas novas
            existing_texts = {line.split('\t')[0] for line in existing_lines if '\t' in line}
            new_lines = [line for line in chinese_lines if line not in existing_texts]
            
            if new_lines:
                for line in new_lines:
                    f.write(f"{line}\n")
                print(f"✅ Adicionadas {len(new_lines)} linhas novas ao arquivo base")
            else:
                print("ℹ️  Nenhuma linha nova para adicionar")
    
    return True


def process_base_with_llm(base_path: Path, force_provider: str | None = None) -> bool:
    """
    Processa o arquivo base.txt chamando LLM para gerar pares de tradução.
    
    Args:
        base_path: Caminho para o arquivo base.txt
        force_provider: Força uso de provider específico ('maritaca' ou 'deepseek')
        
    Returns:
        bool: True se sucesso
    """
    print(f"🤖 Processando base.txt com LLM...")
    
    try:
        # Lê o arquivo base
        with open(base_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # Processa cada linha
        processed_lines = []
        pairs_cache = {}
        
        for i, line in enumerate(lines, 1):
            line = line.strip()
            if not line:
                continue
            
            # Se a linha já tem pares (formato: texto\tpares_json), pula
            if '\t' in line:
                processed_lines.append(line)
                continue
            
            print(f"   🔄 Processando linha {i}/{len(lines)}: {line[:50]}...")
            
            # Gera pares usando LLM
            try:
                provider = _get_api_provider(force_provider)
                if provider == "maritaca":
                    pairs_str = _retry_api_call(_call_maritaca_pairs, line)
                else:
                    pairs_str = _retry_api_call(_call_deepseek_pairs, line)
                
                # Adiciona linha com pares
                processed_lines.append(f"{line}\t{pairs_str}")
                
                # Pequena pausa entre chamadas
                time.sleep(0.1)
                
            except Exception as e:
                print(f"   ❌ Erro ao processar linha {i}: {e}")
                # Adiciona linha sem pares em caso de erro
                processed_lines.append(f"{line}\t[]")
        
        # Salva arquivo processado
        with open(base_path, 'w', encoding='utf-8') as f:
            for line in processed_lines:
                f.write(f"{line}\n")
        
        print(f"✅ Base.txt processado com LLM")
        return True
        
    except Exception as e:
        print(f"❌ Erro ao processar base.txt com LLM: {e}")
        return False


def sanitize_base_with_word_api(base_path: Path) -> bool:
    """
    Sanitiza o arquivo base.txt usando word-api.
    
    Args:
        base_path: Caminho para o arquivo base.txt
        
    Returns:
        bool: True se sucesso
    """
    print(f"🧹 Sanitizando base.txt com word-api...")
    
    try:
        # Lê o arquivo base
        with open(base_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        processed_lines = []
        modified_count = 0
        removed_count = 0
        
        for line_num, line in enumerate(lines, 1):
            line = line.strip()
            if not line:
                continue
            
            # Divide a linha por tabs
            parts = line.split('\t')
            if len(parts) < 2:
                processed_lines.append(line)
                continue
            
            text = parts[0]
            pairs_str = parts[1]
            
            # Processa pares com word-api
            if pairs_str and pairs_str != "[]":
                pairs = extract_pairs_from_translation(pairs_str)
                if pairs:
                    filtered_pairs = process_word_api_integration(pairs)
                    
                    if filtered_pairs:
                        # Reconstrói pares filtrados
                        new_pairs_parts = []
                        for pair in filtered_pairs:
                            if pair["pinyin"]:
                                new_pairs_parts.append(f'"{pair["word"]} ({pair["pinyin"]}): {pair["translation"]}"')
                            else:
                                new_pairs_parts.append(f'"{pair["word"]}: {pair["translation"]}"')
                        
                        new_pairs_str = "[" + ", ".join(new_pairs_parts) + "]"
                        parts[1] = new_pairs_str
                        
                        if new_pairs_str != pairs_str:
                            modified_count += 1
                    else:
                        # Todos os pares foram removidos, pula a linha
                        removed_count += 1
                        continue
            
            processed_lines.append('\t'.join(parts))
        
        # Salva arquivo sanitizado
        with open(base_path, 'w', encoding='utf-8') as f:
            for line in processed_lines:
                f.write(f"{line}\n")
        
        print(f"✅ Base.txt sanitizado com word-api")
        print(f"   📝 {modified_count} linhas modificadas")
        print(f"   🗑️  {removed_count} linhas removidas")
        return True
        
    except Exception as e:
        print(f"❌ Erro ao sanitizar base.txt: {e}")
        return False


def generate_docx_with_translations(pdf_path: Path, base_path: Path, output_path: Path) -> bool:
    """
    Gera novo DOCX com pinyin e traduções em formato de torre.
    
    Args:
        pdf_path: Caminho para o PDF original
        base_path: Caminho para o arquivo base.txt
        output_path: Caminho para o DOCX de saída
        
    Returns:
        bool: True se sucesso
    """
    print(f"📄 Gerando DOCX com traduções...")
    
    try:
        # Tenta usar python-docx primeiro
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            return _generate_docx_with_python_docx(pdf_path, base_path, output_path)
            
        except ImportError:
            print("⚠️  python-docx não encontrado, tentando PyMuPDF...")
            
            # Fallback para PyMuPDF
            try:
                import fitz  # PyMuPDF
                return _generate_pdf_with_pymupdf(pdf_path, base_path, output_path)
                
            except ImportError:
                print("⚠️  PyMuPDF não encontrado, gerando arquivo de texto...")
                return _generate_text_with_translations(pdf_path, base_path, output_path)
            
    except Exception as e:
        print(f"❌ Erro ao gerar DOCX com traduções: {e}")
        return False


def _generate_docx_with_python_docx(pdf_path: Path, base_path: Path, output_path: Path) -> bool:
    """Gera DOCX usando python-docx com torre pinyin | caracter | tradução."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    
    # Lê o arquivo base
    with open(base_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Cria mapeamento de texto original para traduções
    text_to_translations = {}
    for line in lines:
        line = line.strip()
        if not line or '\t' not in line:
            continue
        
        parts = line.split('\t', 1)
        if len(parts) == 2:
            text = parts[0]
            pairs_str = parts[1]
            
            # Extrai pares de tradução
            pairs = extract_pairs_from_translation(pairs_str)
            if pairs:
                text_to_translations[text] = pairs
    
    # Cria novo documento DOCX
    doc = Document()
    
    # Configuração da página
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Processa cada linha do PDF original
    pdf_text = extract_text_from_pdf(pdf_path)
    
    for line in pdf_text:
        line = line.strip()
        if not line:
            continue
        
        # Verifica se contém caracteres chineses
        if contains_chinese_characters(line):
            # Procura traduções para esta linha
            if line in text_to_translations:
                pairs = text_to_translations[line]
                
                # Cria parágrafo para a torre
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Adiciona cada palavra da torre
                for i, pair in enumerate(pairs):
                    chinese = pair["word"]
                    pinyin = pair["pinyin"]
                    translation = pair["translation"]
                    
                    if chinese:
                        # Cria tabela de 3 linhas para cada palavra
                        table = doc.add_table(rows=3, cols=1)
                        table.style = 'Table Grid'
                        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Pinyin (linha superior)
                        if pinyin:
                            pinyin_cell = table.cell(0, 0)
                            pinyin_cell.text = pinyin
                            pinyin_para = pinyin_cell.paragraphs[0]
                            pinyin_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            pinyin_run = pinyin_para.runs[0]
                            pinyin_run.font.size = Pt(10)
                            pinyin_run.font.color.rgb = RGBColor(147, 112, 219)  # Roxo
                            pinyin_run.font.bold = True
                        
                        # Chinês (linha do meio)
                        chinese_cell = table.cell(1, 0)
                        chinese_cell.text = chinese
                        chinese_para = chinese_cell.paragraphs[0]
                        chinese_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        chinese_run = chinese_para.runs[0]
                        chinese_run.font.size = Pt(12)
                        chinese_run.font.bold = True
                        
                        # Tradução (linha inferior)
                        if translation:
                            translation_cell = table.cell(2, 0)
                            translation_cell.text = translation
                            translation_para = translation_cell.paragraphs[0]
                            translation_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            translation_run = translation_para.runs[0]
                            translation_run.font.size = Pt(9)
                            translation_run.font.color.rgb = RGBColor(255, 140, 0)  # Laranja
                        
                        # Espaço entre palavras
                        if i < len(pairs) - 1:
                            doc.add_paragraph(" ")  # Espaço entre palavras
                
                # Linha em branco após cada frase
                doc.add_paragraph()
            else:
                # Linha sem tradução - mantém original
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Linha sem caracteres chineses - mantém original
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Salva o documento
    doc.save(output_path)
    
    print(f"✅ DOCX com torres gerado: {output_path.name}")
    return True


def _generate_pdf_with_pymupdf(pdf_path: Path, base_path: Path, output_path: Path) -> bool:
    """Gera PDF usando PyMuPDF com torre pinyin | caracter | tradução."""
    import fitz  # PyMuPDF
    
    # Lê o arquivo base
    with open(base_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Cria mapeamento de texto original para traduções
    text_to_translations = {}
    for line in lines:
        line = line.strip()
        if not line or '\t' not in line:
            continue
        
        parts = line.split('\t', 1)
        if len(parts) == 2:
            text = parts[0]
            pairs_str = parts[1]
            
            # Extrai pares de tradução
            pairs = extract_pairs_from_translation(pairs_str)
            if pairs:
                text_to_translations[text] = pairs
    
    # Abre o PDF original
    doc = fitz.open(pdf_path)
    
    # Processa cada página
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        # Extrai texto da página para encontrar posições
        text_dict = page.get_text("dict")
        
        # Encontra e substitui texto chinês por torres
        for block in text_dict["blocks"]:
            if "lines" in block:
                for line in block["lines"]:
                    for span in line["spans"]:
                        text = span["text"]
                        
                        # Verifica se contém caracteres chineses
                        if contains_chinese_characters(text):
                            # Procura traduções para este texto
                            if text in text_to_translations:
                                pairs = text_to_translations[text]
                                
                                # Cria a torre pinyin | caracter | tradução
                                tower_text = _create_tower_text(pairs)
                                
                                if tower_text:
                                    # Obtém a posição do texto original
                                    bbox = span["bbox"]  # (x0, y0, x1, y1)
                                    
                                    # Remove o texto original
                                    page.add_redact_annot(bbox, fill=(1, 1, 1))  # Fundo branco
                                    page.apply_redactions()
                                    
                                    # Adiciona a torre na mesma posição
                                    _add_tower_to_page(page, tower_text, bbox, pairs)
    
    # Salva o PDF modificado
    doc.save(output_path)
    doc.close()
    
    print(f"✅ PDF com torres gerado: {output_path.name}")
    return True


def _create_tower_text(pairs: list) -> str:
    """Cria texto da torre pinyin | caracter | tradução."""
    if not pairs:
        return ""
    
    # Agrupa caracteres em palavras e cria torres
    tower_lines = []
    
    for pair in pairs:
        chinese = pair["word"]
        pinyin = pair["pinyin"]
        translation = pair["translation"]
        
        if chinese:
            # Cria linha da torre: pinyin | caracter | tradução
            tower_line = f"{pinyin} | {chinese} | {translation}"
            tower_lines.append(tower_line)
    
    return "\n".join(tower_lines)


def _add_tower_to_page(page, tower_text: str, bbox: tuple, pairs: list):
    """Adiciona torre de texto à página do PDF."""
    import fitz
    
    # Configurações de fonte
    font_size = 12
    line_height = font_size * 1.3
    word_spacing = font_size * 1.5  # Espaçamento entre palavras
    
    # Posição inicial (canto superior esquerdo do bbox original)
    x0, y0, x1, y1 = bbox
    start_x = x0
    start_y = y0
    
    # Cores (similar ao process_chunks.py)
    pinyin_color = (0.58, 0.44, 0.86)  # #9370DB (roxo)
    chinese_color = (0, 0, 0)  # preto (melhor visibilidade)
    translation_color = (0.8, 0.4, 0)  # laranja (melhor que amarelo)
    
    # Calcula altura total da torre para centralizar verticalmente
    tower_height = (font_size * 0.8) + line_height + (font_size * 0.7) + (line_height * 0.5)
    center_y = y0 + (y1 - y0) / 2
    tower_start_y = center_y - tower_height / 2
    
    # Adiciona cada palavra da torre horizontalmente
    current_x = start_x
    for i, pair in enumerate(pairs):
        chinese = pair["word"]
        pinyin = pair["pinyin"]
        translation = pair["translation"]
        
        if chinese:
            # Calcula largura da palavra para centralizar
            word_width = max(len(chinese) * font_size * 0.6, len(pinyin) * font_size * 0.5, len(translation) * font_size * 0.4)
            word_center_x = current_x + word_width / 2
            
            # Pinyin (linha superior) - centralizado
            if pinyin:
                pinyin_x = word_center_x - (len(pinyin) * font_size * 0.25)
                page.insert_text(
                    (pinyin_x, tower_start_y),
                    pinyin,
                    fontsize=font_size * 0.8,
                    color=pinyin_color
                )
            
            # Chinês (linha do meio) - centralizado
            chinese_x = word_center_x - (len(chinese) * font_size * 0.3)
            page.insert_text(
                (chinese_x, tower_start_y + line_height),
                chinese,
                fontsize=font_size,
                color=chinese_color
            )
            
            # Tradução (linha inferior) - centralizado
            if translation:
                translation_x = word_center_x - (len(translation) * font_size * 0.2)
                page.insert_text(
                    (translation_x, tower_start_y + line_height * 2),
                    translation,
                    fontsize=font_size * 0.7,
                    color=translation_color
                )
            
            # Move para próxima palavra
            current_x += word_width + word_spacing


def _generate_pdf_with_reportlab(pdf_path: Path, base_path: Path, output_path: Path) -> bool:
    """Gera PDF usando ReportLab."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    
    # Lê o arquivo base
    with open(base_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Cria mapeamento de texto original para traduções
    text_to_translations = {}
    for line in lines:
        line = line.strip()
        if not line or '\t' not in line:
            continue
        
        parts = line.split('\t', 1)
        if len(parts) == 2:
            text = parts[0]
            pairs_str = parts[1]
            
            # Extrai pares de tradução
            pairs = extract_pairs_from_translation(pairs_str)
            if pairs:
                text_to_translations[text] = pairs
    
    # Cria novo PDF
    output_canvas = canvas.Canvas(str(output_path), pagesize=letter)
    
    y_position = 750  # Posição inicial Y
    line_height = 20
    
    for text, pairs in text_to_translations.items():
        # Desenha pinyin
        pinyin_text = " ".join([pair["pinyin"] for pair in pairs if pair["pinyin"]])
        if pinyin_text:
            output_canvas.setFont("Helvetica", 10)
            output_canvas.drawString(50, y_position, pinyin_text)
            y_position -= line_height
        
        # Desenha texto chinês
        output_canvas.setFont("Helvetica", 12)
        output_canvas.drawString(50, y_position, text)
        y_position -= line_height
        
        # Desenha traduções
        translation_text = " ".join([pair["translation"] for pair in pairs])
        if translation_text:
            output_canvas.setFont("Helvetica", 10)
            output_canvas.drawString(50, y_position, translation_text)
            y_position -= line_height
        
        y_position -= 10  # Espaço extra entre grupos
        
        # Verifica se precisa de nova página
        if y_position < 50:
            output_canvas.showPage()
            y_position = 750
    
    output_canvas.save()
    print(f"✅ PDF com traduções gerado: {output_path.name}")
    return True


def _generate_text_with_translations(pdf_path: Path, base_path: Path, output_path: Path) -> bool:
    """Gera arquivo de texto com traduções (fallback)."""
    # Lê o arquivo base
    with open(base_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Gera arquivo de texto
    text_output_path = output_path.with_suffix('.txt')
    
    with open(text_output_path, 'w', encoding='utf-8') as f:
        f.write("PDF com Traduções - Texto Chinês, Pinyin e Português\n")
        f.write("=" * 60 + "\n\n")
        
        for line in lines:
            line = line.strip()
            if not line or '\t' not in line:
                continue
            
            parts = line.split('\t', 1)
            if len(parts) == 2:
                text = parts[0]
                pairs_str = parts[1]
                
                # Extrai pares de tradução
                pairs = extract_pairs_from_translation(pairs_str)
                if pairs:
                    f.write(f"Texto chinês: {text}\n")
                    
                    # Pinyin
                    pinyin_text = " ".join([pair["pinyin"] for pair in pairs if pair["pinyin"]])
                    if pinyin_text:
                        f.write(f"Pinyin: {pinyin_text}\n")
                    
                    # Traduções
                    translation_text = " ".join([pair["translation"] for pair in pairs])
                    if translation_text:
                        f.write(f"Português: {translation_text}\n")
                    
                    f.write("-" * 40 + "\n\n")
    
    print(f"✅ Arquivo de texto com traduções gerado: {text_output_path.name}")
    print("   (Instale reportlab para gerar PDF: pip install reportlab)")
    return True


def main():
    parser = argparse.ArgumentParser(
        description="Processa PDFs com texto chinês e gera versão com pinyin e traduções",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemplos:
  python3 text_burner.py texto    # Processa assets/texto/

Funcionamento:
  - Procura por arquivo PDF no diretório assets/<directory_name>
  - Extrai linhas contendo caracteres chineses
  - Cria base.txt com essas linhas
  - Chama LLM para gerar pares de tradução
  - Sanitiza usando word-api
  - Gera novo PDF com pinyin e traduções sobrepostas
  - É idempotente - continua de onde parou se base.txt já existir
        """
    )
    
    parser.add_argument('directory', help='Nome do diretório dentro de assets/')
    
    # API provider selection
    api_group = parser.add_mutually_exclusive_group()
    api_group.add_argument(
        "-m", "--maritaca",
        action="store_true",
        help="Force use of Maritaca AI API (requires MARITACA_API_KEY)"
    )
    api_group.add_argument(
        "-d", "--deepseek",
        action="store_true",
        help="Force use of DeepSeek API (requires DEEPSEEK_API_KEY)"
    )
    
    args = parser.parse_args()
    
    # Determina provider forçado
    force_provider = None
    if args.maritaca:
        force_provider = "maritaca"
    elif args.deepseek:
        force_provider = "deepseek"
    
    # Constrói caminhos
    assets_root = Path(__file__).resolve().parent / "assets"
    source_dir = assets_root / args.directory
    
    print("🔥 Text Burner - Processador de PDFs com texto chinês")
    print("=" * 60)
    print(f"📁 Diretório: {source_dir}")
    
    # Verifica se o diretório existe
    if not source_dir.exists():
        print(f"❌ Erro: Diretório {source_dir} não encontrado")
        return 1
    
    # Procura por arquivo PDF
    pdf_files = list(source_dir.glob("*.pdf"))
    if not pdf_files:
        print(f"❌ Erro: Nenhum arquivo PDF encontrado em {source_dir}")
        return 1
    
    if len(pdf_files) > 1:
        print(f"⚠️  Múltiplos PDFs encontrados, usando: {pdf_files[0].name}")
    
    pdf_path = pdf_files[0]
    base_path = source_dir / "base.txt"
    output_path = source_dir / f"{pdf_path.stem}_with_translations.docx"
    
    print(f"📄 PDF encontrado: {pdf_path.name}")
    
    # 1. Cria arquivo base.txt
    print("\n📝 Passo 1: Criando base.txt...")
    if not create_base_file(pdf_path, base_path, resume=True):
        return 1
    
    # 2. Processa com LLM
    print("\n🤖 Passo 2: Processando com LLM...")
    if not process_base_with_llm(base_path, force_provider):
        return 1
    
    # 3. Verifica word-api e sanitiza
    print("\n🧹 Passo 3: Sanitizando com word-api...")
    if not check_word_api_health():
        print("⚠️  Word-api indisponível, pulando sanitização")
    else:
        if not sanitize_base_with_word_api(base_path):
            return 1
    
    # 4. Gera DOCX com traduções
    print("\n📄 Passo 4: Gerando DOCX com traduções...")
    if not generate_docx_with_translations(pdf_path, base_path, output_path):
        return 1
    
    print("\n🎉 Processamento concluído com sucesso!")
    print(f"📄 DOCX com traduções: {output_path.name}")
    print(f"📝 Arquivo base: {base_path.name}")
    
    return 0


if __name__ == "__main__":
    sys.exit(main())
